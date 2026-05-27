from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
import os

BASE = os.path.dirname(__file__)
IMAGES_DIR = os.path.join(BASE, '_images')
PLACEHOLDER_DIR = os.path.join(BASE, '_placeholder-images')
OUT_DOC = os.path.join(BASE, 'Emerald_City_Rounds_Style_Guide_GENERATED.docx')

palette = [
    ("Background Off-White", "#f9fbff"),
    ("Primary Blue (Header)", "#0966C2"),
    ("Navigation Gray", "#a6a6a6"),
    ("Accent Green (Hover)", "#09c20f"),
    ("Accent Gold (Icon Hover)", "#F3BA64"),
    ("Section Teal (Cards)", "#0d93a7"),
    ("Footer Teal", "#008080"),
]

# helper to hex -> rgb
def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

# create swatch images
SWATCH_DIR = os.path.join(BASE, 'styleguide_swatches')
os.makedirs(SWATCH_DIR, exist_ok=True)
swatch_paths = []
for name, hexv in palette:
    rgb = hex_to_rgb(hexv)
    img = Image.new('RGB', (400, 200), rgb)
    draw = ImageDraw.Draw(img)
    # write label on swatch (white or black depending on brightness)
    brightness = (rgb[0]*299 + rgb[1]*587 + rgb[2]*114) / 1000
    text_color = (0,0,0) if brightness > 130 else (255,255,255)
    try:
        font = ImageFont.load_default()
    except Exception:
        font = None
    draw.text((10,10), f"{name} {hexv} RGB{rgb}", fill=text_color, font=font)
    path = os.path.join(SWATCH_DIR, f"swatch_{hexv.lstrip('#')}.png")
    img.save(path)
    swatch_paths.append((name, hexv, rgb, path))

# Compose the docx
doc = Document()

doc.add_heading('Emerald City Rounds — Site Style Guide', level=1)

# Project summary (inferred/updated from site files)
doc.add_heading('Project Overview', level=2)
doc.add_paragraph('Project Name: Emerald City Rounds')
doc.add_paragraph('Audience: Local community, students, and participants interested in music and community events.')
doc.add_paragraph('Purpose: Promote classes, events, and resources; provide information and contact for Emerald City Rounds.')
doc.add_paragraph('Personality: Friendly, community-focused, approachable, and slightly playful while remaining professional.')
doc.add_paragraph('Comparable sites: Local dance or music studios, community arts centers, Cascadia College events pages.')
doc.add_paragraph('Features/Content: Home banner carousel, Events, Calendar, Classes, Resources, Account/Login, Store, Footer with social links.')

# Color palette
doc.add_heading('Color Palette', level=2)
for name, hexv, rgb, path in swatch_paths:
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(path, width=Inches(2))
    p.add_run(f'  {name} — HEX: {hexv} — RGB: {rgb}')

# Typography and overall styling
doc.add_heading('Typography & Overall Styling', level=2)
doc.add_paragraph('Base font: Arial, sans-serif. Font-size base: 16px. Line-height: 1.5. Layout uses a flexible column flow with full-bleed header/banner images; site uses responsive flexbox-based grids for content areas.')

# Section styles
doc.add_heading('Header', level=2)
doc.add_paragraph('Style: Full-width banner (full-bleed) image or colored header bar with centered content. Background color: #0966C2. Text color: white. Padding: approx 1rem. H1 size approx 2rem, bold.')

doc.add_heading('Navigation', level=2)
doc.add_paragraph('Style: Centered horizontal navigation with large tap targets. Background color: #a6a6a6; border-top/bottom: 3px solid #0966C2. Links: white text, hover: background #09c20f and icon color change to #F3BA64. Font-size approx 1.25–2rem.')

doc.add_heading('Main / Content Areas', level=2)
doc.add_paragraph('Style: Flexible two-column layout (50%/50% blocks for cards). Card/background color: #0d93a7 with white text for some feature panels. Body background: #f9fbff. Use generous padding and responsive stacking for narrow screens.')

doc.add_heading('Footer', level=2)
doc.add_paragraph('Style: Full-width footer with background #008080 and white text. Includes multi-column footer layout (three columns) and social icons. Padding ~15px. Text aligned left.')

# Layout sketch placeholder
doc.add_heading('Layout Sketch', level=2)
doc.add_paragraph('Include a hand-drawn layout sketch photo here. Replace the placeholder image with your photographed sketch of the final site layout.')
# Use an existing image as placeholder if present
placeholder_img = None
if os.path.isdir(IMAGES_DIR):
    for fname in os.listdir(IMAGES_DIR):
        if fname.lower().endswith(('.jpg', '.jpeg', '.png')):
            placeholder_img = os.path.join(IMAGES_DIR, fname)
            break
if placeholder_img:
    doc.add_picture(placeholder_img, width=Inches(6))
else:
    doc.add_paragraph('[PLACEHOLDER: add photo of hand-drawn layout here]')

# Media licensing
doc.add_heading('Media & Licensing', level=2)
doc.add_paragraph('This table lists media files used in the site and a recommended licensing note. Verify licenses and replace notes with accurate links when available.')

media_files = []
if os.path.isdir(IMAGES_DIR):
    for fname in sorted(os.listdir(IMAGES_DIR)):
        media_files.append((os.path.join('_images', fname), 'Author-created — assumed copyright owner: Leo Urquhart. Replace with a license link if third-party.'))
if os.path.isdir(PLACEHOLDER_DIR):
    for fname in sorted(os.listdir(PLACEHOLDER_DIR)):
        media_files.append((os.path.join('_placeholder-images', fname), 'Placeholder image included with project — verify source and license before reuse.'))

# Add simple table-like listing
for path, note in media_files:
    p = doc.add_paragraph()
    p.add_run(path + ': ').bold = True
    p.add_run(note)

# Brand/social icons guidance
doc.add_heading('Social Icons / Brand Assets', level=2)
doc.add_paragraph('Social icons included in _placeholder-images or _images (Twitter, Facebook, Instagram, YouTube, LinkedIn, Medium). Use each platform\'s brand resources and follow their guidelines:')
doc.add_paragraph('Twitter: https://about.twitter.com/en/who-we-are/brand-toolkit')
doc.add_paragraph('Facebook: https://en.facebookbrand.com/')
doc.add_paragraph('Instagram: https://about.instagram.com/brand')
doc.add_paragraph('YouTube: https://www.youtube.com/about/brand-resources/')

# Save the document
print('Saving to', OUT_DOC)
doc.save(OUT_DOC)
print('Done')
